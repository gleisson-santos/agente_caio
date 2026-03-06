#!/usr/bin/env python3
"""
Gerador de documentos Word (.docx)
Uso: python generate_docx.py [arquivo_saida.docx]

Dependências: pip install python-docx
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
import sys
import os

def criar_documento(caminho_saida="out/documento.docx"):
    """Cria um documento Word de exemplo"""

    doc = Document()

    # === CONFIGURAÇÕES DE PÁGINA ===
    section = doc.sections[0]
    section.page_width = Cm(21)  # A4
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # === TÍTULO ===
    titulo = doc.add_heading('Título do Documento', level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # === PARÁGRAFO ===
    p = doc.add_paragraph()
    p.add_run('Este é um documento de exemplo ').bold = False
    p.add_run('com texto em negrito').bold = True
    p.add_run(' e ')
    p.add_run('texto em itálico').italic = True
    p.add_run('.')

    # === SUBTÍTULO ===
    doc.add_heading('Seção 1: Introdução', level=1)
    doc.add_paragraph(
        'Lorem ipsum dolor sit amet, consectetur adipiscing elit. '
        'Sed do eiusmod tempor incididunt ut labore et dolore magna aliqua.'
    )

    # === LISTA COM MARCADORES ===
    doc.add_heading('Lista de Itens', level=2)
    for item in ['Primeiro item', 'Segundo item', 'Terceiro item']:
        doc.add_paragraph(item, style='List Bullet')

    # === LISTA NUMERADA ===
    doc.add_heading('Lista Numerada', level=2)
    for item in ['Passo um', 'Passo dois', 'Passo três']:
        doc.add_paragraph(item, style='List Number')

    # === TABELA ===
    doc.add_heading('Tabela de Dados', level=1)

    dados = [
        ['Nome', 'Idade', 'Cidade'],
        ['Ana', '25', 'São Paulo'],
        ['Bruno', '30', 'Rio de Janeiro'],
        ['Carla', '28', 'Belo Horizonte'],
    ]

    tabela = doc.add_table(rows=len(dados), cols=len(dados[0]))
    tabela.style = 'Table Grid'
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Preencher tabela
    for i, linha in enumerate(dados):
        for j, celula in enumerate(linha):
            tabela.rows[i].cells[j].text = celula
            # Cabeçalho em negrito
            if i == 0:
                tabela.rows[i].cells[j].paragraphs[0].runs[0].bold = True

    # === CITAÇÃO ===
    doc.add_heading('Citação', level=2)
    citacao = doc.add_paragraph(
        '"A imaginação é mais importante que o conhecimento." - Albert Einstein',
        style='Quote'
    )

    # === QUEBRA DE PÁGINA ===
    doc.add_page_break()

    # === NOVA PÁGINA ===
    doc.add_heading('Seção 2: Conclusão', level=1)
    doc.add_paragraph(
        'Este documento foi gerado automaticamente usando python-docx. '
        'Você pode modificar este script para criar documentos personalizados.'
    )

    # === RODAPÉ (nota) ===
    doc.add_paragraph()
    rodape = doc.add_paragraph()
    rodape.add_run('Documento gerado automaticamente').italic = True
    rodape.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Criar diretório se não existir
    os.makedirs(os.path.dirname(caminho_saida) if os.path.dirname(caminho_saida) else '.', exist_ok=True)

    # Salvar
    doc.save(caminho_saida)
    print(f"Documento criado: {caminho_saida}")
    return caminho_saida


def criar_documento_personalizado(titulo, secoes, caminho_saida="out/documento.docx"):
    """
    Cria documento personalizado

    Args:
        titulo: Título do documento
        secoes: Lista de dicts com {'titulo': str, 'conteudo': str}
        caminho_saida: Caminho do arquivo de saída
    """
    doc = Document()

    # Título
    t = doc.add_heading(titulo, level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Seções
    for secao in secoes:
        doc.add_heading(secao.get('titulo', 'Seção'), level=1)
        doc.add_paragraph(secao.get('conteudo', ''))

    os.makedirs(os.path.dirname(caminho_saida) if os.path.dirname(caminho_saida) else '.', exist_ok=True)
    doc.save(caminho_saida)
    print(f"Documento criado: {caminho_saida}")
    return caminho_saida


# === FUNÇÕES AUXILIARES PARA USO COM LLMs ===

def adicionar_titulo(doc, texto, nivel=0):
    """Adiciona título ao documento"""
    return doc.add_heading(texto, level=nivel)

def adicionar_paragrafo(doc, texto, negrito=False, italico=False):
    """Adiciona parágrafo ao documento"""
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.bold = negrito
    run.italic = italico
    return p

def adicionar_tabela(doc, dados, estilo='Table Grid'):
    """Adiciona tabela ao documento (dados = lista de listas)"""
    if not dados:
        return None
    tabela = doc.add_table(rows=len(dados), cols=len(dados[0]))
    tabela.style = estilo
    for i, linha in enumerate(dados):
        for j, celula in enumerate(linha):
            tabela.rows[i].cells[j].text = str(celula)
    return tabela

def adicionar_lista(doc, itens, numerada=False):
    """Adiciona lista ao documento"""
    estilo = 'List Number' if numerada else 'List Bullet'
    for item in itens:
        doc.add_paragraph(item, style=estilo)


if __name__ == "__main__":
    saida = sys.argv[1] if len(sys.argv) > 1 else "out/documento_exemplo.docx"
    criar_documento(saida)
